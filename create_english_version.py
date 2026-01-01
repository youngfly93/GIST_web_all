#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建ChatGIST项目技术实现详解的英文版本
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_english_version():
    """创建英文版本的技术文档"""
    
    # 文档路径
    source_path = "ChatGIST项目技术实现详解_before_enhancement.docx"
    
    if not os.path.exists(source_path):
        print(f"错误：找不到源文档文件 {source_path}")
        return False
    
    try:
        # 打开源文档
        print(f"正在读取源文档: {source_path}")
        
        # 创建新的英文文档
        doc = Document()
        
        # 添加英文版本内容
        english_content = [
            ("# ChatGIST Project Technical Implementation Details", "title"),
            ("", "normal"),
            ("## Overview", "heading"),
            ("", "normal"),
            ("ChatGIST is a comprehensive bioinformatics analysis platform focused on Gastrointestinal Stromal Tumor (GIST) research, integrating modern web technologies and artificial intelligence capabilities. The project adopts a hybrid multi-stack architecture, including React frontend, Node.js backend, and three independent R Shiny analysis modules, providing a complete solution from gene queries to in-depth data analysis for biomedical researchers.", "normal"),
            ("", "normal"),
            ("The core feature of the project lies in combining traditional bioinformatics analysis tools with modern AI technology, significantly lowering the barrier to using professional analysis tools through intelligent user interfaces and automated data interpretation. Meanwhile, the project employs modular design and microservice architecture, ensuring system scalability and maintainability, laying a solid foundation for future functional expansion.", "normal"),
            ("", "normal"),
            ("## Part 1: ChatGIST System Architecture", "heading"),
            ("", "normal"),
            ("### 1.1 Hybrid Multi-Stack Architecture Design", "subheading"),
            ("", "normal"),
            ("ChatGIST adopts an innovative hybrid multi-stack architecture, organically combining modern web technologies with professional bioinformatics analysis tools. The entire system consists of three main layers: React frontend for user interaction and data display, Node.js backend for API request handling and AI integration, and R Shiny modules for professional bioinformatics analysis functions.", "normal"),
            ("", "normal"),
            ("The core design philosophy is modularity and scalability. The frontend uses React + TypeScript + Vite modern technology stack, ensuring development efficiency and code quality. The backend uses Express framework to build RESTful APIs, providing unified data interfaces and AI service proxies. The three R Shiny analysis modules are deployed independently on different ports, accessed through nginx reverse proxy, ensuring both system stability and ease of independent maintenance and expansion.", "normal"),
            ("", "normal"),
            ("### 1.2 Frontend Technical Implementation", "subheading"),
            ("", "normal"),
            ("ChatGIST's frontend adopts React 18 + TypeScript + Vite modern technology stack, providing responsive user interfaces and smooth interaction experiences. The project structure is clear, including seven main pages: Home, GeneInfo, AIChat, MiRNAResults, Dataset, Guide, and GistDatabase. Each page is developed using functional components and Hooks, ensuring code conciseness and maintainability.", "normal"),
            ("", "normal"),
            ("In terms of component design, the system implements several core components, including FloatingChat, GeneAssistant, SmartCapture, and PageNavigator. These components adopt modern React design patterns, supporting property passing, state management, and event handling. Particularly noteworthy is the SmartCapture component, which implements page screenshot functionality, allowing screenshots to be sent directly to AI for analysis, greatly enhancing user experience.", "normal"),
            ("", "normal"),
            ("Route management uses React Router v6, implementing single-page application navigation. State management employs React's built-in useState and useContext, with useReducer for complex state logic. The styling system combines CSS modules and inline styles, ensuring style isolation and maintainability. The build tool Vite provides fast development servers and efficient production builds, significantly improving development experience.", "normal"),
            ("", "normal"),
            ("### 1.3 Backend API Architecture", "subheading"),
            ("", "normal"),
            ("The backend uses Node.js + Express framework, providing four main API routes: /api/chat (AI conversation), /api/gene (gene information), /api/ncrna (non-coding RNA), and /api/proxy (proxy service). Each route has clear responsibility division: chat route handles AI conversation requests, gene route provides gene information queries, ncrna route processes non-coding RNA related queries, and proxy route serves as a proxy for external APIs.", "normal"),
            ("", "normal"),
            ("AI integration is one of the core backend functions, mainly implemented through the Volcano Engine API. The system supports multimodal AI analysis of text and images, processing user text questions and image uploads. API calls adopt both streaming and non-streaming modes: streaming mode provides real-time response experience, while non-streaming mode ensures data integrity. The error handling mechanism is comprehensive, including API key validation, request timeout handling, and friendly prompts for exceptional situations.", "normal"),
            ("", "normal"),
            ("In terms of data services, the system implements two core services: geneFetcher and ncRNAService. geneFetcher is responsible for retrieving gene information from external databases like PubChem, supporting gene symbol queries and detailed information retrieval. ncRNAService handles non-coding RNA related data queries, including miRNA, circRNA, and lncRNA interaction information. These services employ caching mechanisms and error retry strategies, ensuring data retrieval reliability and performance.", "normal"),
            ("", "normal"),
            ("### 1.4 Database and Storage Design", "subheading"),
            ("", "normal"),
            ("ChatGIST's data storage adopts a hybrid strategy, combining file storage and memory caching. Bioinformatics data is primarily stored in file form, including gene expression matrices (RData format), interaction network data (TXT format), and clinical information (CSV format). This design fully considers the characteristics of bioinformatics data, ensuring both data integrity and efficient processing by R language.", "normal"),
            ("", "normal"),
            ("Frontend data management combines local storage and session storage. User query history and preference settings are stored in localStorage, while temporary analysis results and state information are stored in sessionStorage. This design ensures both continuity of user experience and avoids long-term storage of sensitive data. The system also implements automatic data cleanup mechanisms, periodically clearing expired cache data.", "normal"),
            ("", "normal"),
            ("R Shiny modules' data storage is more specialized, mainly containing dbGIST_matrix.Rdata (gene expression matrix), dbGIST_ImmuneCell.RData (immune cell data), dbGIST_msigdb.RData (MSigDB pathway data), and other professional datasets. These data undergo strict quality control and standardization processing, ensuring the accuracy and reproducibility of analysis results. Data loading adopts a lazy loading strategy, loading relevant datasets only when needed, improving system response speed.", "normal"),
            ("", "normal"),
            ("## Part 2: ChatGIST AI Functionality Implementation", "heading"),
            ("", "normal"),
            ("### 2.1 Multimodal AI Integration Architecture", "subheading"),
            ("", "normal"),
            ("ChatGIST's AI functionality adopts an advanced multimodal integration architecture, supporting intelligent analysis of text and images. The system mainly implements AI functions through the Volcano Engine API, which is based on the DeepSeek-V3 model, possessing powerful natural language understanding and image analysis capabilities. AI integration adopts a proxy mode, with the backend serving as an intermediate layer processing frontend requests, calling external AI services, and processing and optimizing results, ensuring system security and stability.", "normal"),
            ("", "normal"),
            ("In terms of technical implementation, the system supports both streaming and non-streaming AI response modes. Streaming mode implements real-time responses through Server-Sent Events (SSE), allowing users to see the AI generating answers character by character, providing better interactive experiences. Non-streaming mode returns complete results at once, suitable for scenarios requiring complete data processing. The two modes can be flexibly switched based on user preferences and network conditions, ensuring optimal experiences in different usage scenarios.", "normal"),
            ("", "normal"),
            ("### 2.2 Intelligent Conversation System Implementation", "subheading"),
            ("", "normal"),
            ("ChatGIST's intelligent conversation system is one of the core functions of the entire platform, specifically optimized for the biomedical field. The system processes user conversation requests through the /api/chat route, supporting both pure text conversations and mixed text-image conversations. Conversation history management adopts frontend state management, maintaining complete conversation contexts, ensuring AI can understand users' continuous questions and complex queries.", "normal"),
            ("", "normal"),
            ("In terms of implementation details, the system employs a message queue mechanism to handle concurrent requests, avoiding conflict issues when multiple users access simultaneously. The error handling mechanism is comprehensive, including handling various exceptional situations such as network timeouts, API quotas, and service unavailability. The user interface adopts a modern chat interface design, supporting Markdown rendering, code highlighting, and mathematical formula display, particularly suitable for professional communication in the bioinformatics field.", "normal"),
            ("", "normal"),
            ("### 2.3 Intelligent Image Analysis Functionality", "subheading"),
            ("", "normal"),
            ("ChatGIST's image analysis functionality is one of its unique advantages, specifically optimized for biomedical charts and data visualizations. The system supports multiple image input methods, including direct upload, page screenshot, and drag-and-drop upload. The SmartCapture component implements page screenshot functionality, allowing users to directly capture charts or data from the current page and send them to AI for analysis, greatly simplifying the operation process.", "normal"),
            ("", "normal"),
            ("The image processing workflow includes format conversion, size optimization, and Base64 encoding. The system supports common image formats (PNG, JPEG, GIF, etc.) and automatically adjusts sizes to meet API requirements. AI analysis results include not only descriptions of image content but also professional biological interpretations, including statistical significance analysis, experimental design evaluation, and result interpretation suggestions, providing valuable scientific insights for researchers.", "normal"),
            ("", "normal"),
            ("### 2.4 R Shiny Module AI Enhancement", "subheading"),
            ("", "normal"),
            ("All three R Shiny analysis modules of ChatGIST integrate AI enhancement functionality, implementing intelligent chart analysis through ai_chat_module.R. The system can automatically identify the analysis module the user is currently operating, accurately retrieve corresponding chart data, and call the AI API for in-depth analysis. This integration method maintains R Shiny's original professional analysis capabilities while adding intelligent data interpretation functions.", "normal"),
            ("", "normal"),
            ("The implementation of AI functionality adopts an active module tracking mechanism, recording the user's last operation in each module, ensuring AI analyzes the chart the user is currently focusing on. Analysis results include interpretations of statistical significance, biological meaning explanations, clinical application value assessments, and other dimensions, providing comprehensive professional support for researchers. The system also supports batch analysis and history record queries, facilitating users' in-depth data exploration and comparative analysis.", "normal"),
            ("", "normal"),
            ("## Part 3: ChatGIST Deployment and Operations", "heading"),
            ("", "normal"),
            ("### 3.1 Containerized Deployment Architecture", "subheading"),
            ("", "normal"),
            ("ChatGIST adopts a modern containerized deployment architecture, achieving unified management of multiple services through Docker and Docker Compose. The entire system contains three main containers: gist-web (frontend and backend services), gist-shiny (R Shiny applications), and nginx (reverse proxy). This architectural design ensures both service isolation and security, as well as ease of expansion and maintenance.", "normal"),
            ("", "normal"),
            ("Docker configuration adopts a multi-stage build strategy, using node:18-alpine image for frontend building and dependency installation in the build stage, retaining only necessary production files in the running stage, greatly reducing image size. The gist-shiny container is based on the rocker/shiny:latest image, specifically optimized for R Shiny applications, providing a stable R runtime environment. The nginx container uses the alpine version, configured with SSL support and reverse proxy rules, implementing a unified access entry point.", "normal"),
            ("", "normal"),
            ("### 3.2 Nginx Reverse Proxy Configuration", "subheading"),
            ("", "normal"),
            ("ChatGIST's nginx configuration implements complex routing distribution and load balancing functions. The main configuration file chatgist.conf defines HTTPS redirection, SSL security configuration, and multi-service routing rules. The system supports two deployment modes: production mode (ports 80/443) and development mode (port 81), which can be flexibly switched according to different environment requirements.", "normal"),
            ("", "normal"),
            ("In terms of routing configuration, the root path '/' points to the frontend application, '/api/' path proxies to the backend API service, and '/shiny/' path proxies to the R Shiny application. Particularly noteworthy is that the system also configures three independent Shiny analysis module routes: '/transcriptomics/', '/proteomics/', and '/posttranslational/', corresponding to different analysis functions. WebSocket support ensures normal operation of real-time communication functions, and file upload limits are set to 100MB to support processing of large data files.", "normal"),
            ("", "normal"),
            ("### 3.3 Multi-Port Service Management", "subheading"),
            ("", "normal"),
            ("ChatGIST adopts a multi-port service architecture, with each functional module running on an independent port. The frontend development server runs on ports 5173/5174, the backend API service runs on port 8000, and the three R Shiny analysis modules run on ports 4964/4966 (transcriptomics), 4968/4967 (proteomics), and 4972/4971 (post-translational modifications) respectively. Each module provides both AI version and basic version choices, meeting different user needs.", "normal"),
            ("", "normal"),
            ("Service management adopts automated scripts, including start_all_shiny.sh (start all services), check_shiny_status.sh (check service status), and stop_shiny.sh (stop services). These scripts implement intelligent service management, automatically detecting service status, handling port conflicts, and managing log files. The logging system adopts hierarchical management, with each service having an independent log file, facilitating problem diagnosis and performance monitoring.", "normal"),
            ("", "normal"),
            ("### 3.4 SSL Certificate and Security Configuration", "subheading"),
            ("", "normal"),
            ("ChatGIST implements complete HTTPS security configuration, supporting TLS 1.2 and TLS 1.3 protocols, adopting modern encryption suites to ensure data transmission security. SSL certificate management adopts automated scripts, including certificate application, update, and deployment functions. Security header configuration includes X-Frame-Options, X-Content-Type-Options, and X-XSS-Protection, effectively preventing common web security attacks.", "normal"),
            ("", "normal"),
            ("Environment variable management adopts a layered configuration strategy, with sensitive information like API keys stored in .env files, and different environments using different configuration files. Production environment configuration files are injected through Docker secrets or environment variables, ensuring the security of sensitive information. Meanwhile, the system implements a health check mechanism, monitoring service status through the /health endpoint, facilitating integration with load balancers and monitoring systems.", "normal"),
            ("", "normal"),
            ("## Part 4: Bioinformatics Analysis Module Implementation Cases", "heading"),
            ("", "normal"),
            ("### 4.1 GIST Project Analysis Module Architecture", "subheading"),
            ("", "normal"),
            ("Based on the actual needs of the GIST (Gastrointestinal Stromal Tumor) research project, we have developed three core bioinformatics analysis modules: Transcriptomics, Proteomics, and Post-translational Modification Omics. These modules are built using the R Shiny framework, integrating AI intelligent analysis functions, providing a complete data analysis solution for biomedical research.", "normal"),
            ("", "normal"),
            ("Each analysis module follows unified technical architecture standards, including standardized project structures, consistent theme styles, modular AI chat functionality, and flexible deployment schemes. This standardized approach not only ensures consistency in user experience but also greatly improves development efficiency and system maintainability. The modules adopt independent deployment strategies, providing services through different ports, ensuring both system stability and ease of function expansion and upgrade.", "normal"),
            ("", "normal"),
            ("### 4.2 Transcriptomics Analysis Module", "subheading"),
            ("", "normal"),
            ("The Transcriptomics analysis module is a core component of the GIST project, specifically used for analyzing gene expression data of gastrointestinal stromal tumors. The module is built based on the R Shiny framework, adopting bs4Dash dashboard layout, providing five main analysis functions: Introduction, Module2 (single gene expression analysis), Module3 (gene correlation analysis), Module4 (drug resistance analysis), and Module5 (before and after treatment comparison).", "normal"),
            ("", "normal"),
            ("In terms of technical implementation, the module uses the global.R file to load all necessary R package dependencies, including data processing packages (tidyverse, data.table), visualization packages (ggplot2, ggpubr), statistical analysis packages (pROC), and bioinformatics packages (clusterProfiler, org.Hs.eg.db). Data storage uses RData format, mainly containing gene expression matrices, clinical data, immune cell infiltration data, and pathway database information. The user interface design focuses on interactivity and usability, supporting gene symbol input, parameter selection, and real-time chart updates.", "normal"),
            ("", "normal"),
            ("AI functionality integration is a major highlight of this module, implementing intelligent chart analysis functions through ai_chat_module.R. The system can automatically identify the module the user is currently operating, accurately retrieve corresponding chart data, and call external AI APIs for in-depth analysis. AI analysis results include statistical significance interpretation, biological meaning explanation, and clinical application value assessment, providing professional data interpretation support for researchers.", "normal"),
            ("", "normal"),
            ("### 4.3 Proteomics Analysis Module", "subheading"),
            ("", "normal"),
            ("The Proteomics analysis module focuses on protein-level data analysis, adopting a similar technical architecture to the Transcriptomics module, but specifically optimized for protein data characteristics. The module supports protein expression level analysis, protein interaction network construction, functional enrichment analysis, and differential protein screening, providing comprehensive analysis tools for proteomics research.", "normal"),
            ("", "normal"),
            ("The data processing workflow is optimized for proteomics data characteristics, including missing value handling, data normalization, batch effect correction, and quality control. Visualization functions provide various chart types, such as volcano plots, heatmaps, PCA plots, and protein interaction network diagrams, helping researchers understand proteomics data from different perspectives. Statistical analysis methods cover t-tests, variance analysis, multiple comparison correction, and machine learning algorithms, ensuring the reliability and accuracy of analysis results.", "normal"),
            ("", "normal"),
            ("The module's AI functionality is specifically trained and optimized for proteomics data, capable of identifying protein functional domains, predicting protein interactions, analyzing signaling pathway activity, and interpreting biological processes. AI analysis results include not only statistical interpretations but also protein function annotations, pathway enrichment analysis, and disease association analysis, providing deep biological insights for proteomics research.", "normal"),
            ("", "normal"),
            ("### 4.4 Post-translational Modification Omics Analysis Module", "subheading"),
            ("", "normal"),
            ("The Post-translational Modification Omics analysis module is the most complex of the three modules, specifically processing protein post-translational modification data, especially phosphoproteomics data. The module not only needs to process protein expression data but also analyze modification site information, modification degree changes, and modification functional impacts, providing powerful analysis tools for in-depth understanding of protein function regulation mechanisms.", "normal"),
            ("", "normal"),
            ("In terms of technical implementation, the module integrates multiple professional bioinformatics algorithms and databases, including phosphorylation site prediction, kinase-substrate relationship analysis, signaling pathway activity assessment, and functional enrichment analysis. The data structure design considers the complexity of modification sites, supporting simultaneous analysis of multiple modification types. Visualization functions provide modification site distribution diagrams, kinase activity heatmaps, signaling pathway network diagrams, and time series analysis charts, helping researchers understand modification dynamic change processes.", "normal"),
            ("", "normal"),
            ("AI functionality plays an important role in this module, analyzing modification patterns, predicting functional impacts, and identifying key regulatory nodes through deep learning algorithms. The AI system can integrate multi-layer omics data, providing systematic biological interpretations, including the impact of modifications on protein functions, activation states of signaling pathways, and disease relevance analysis. This intelligent analysis greatly improves research efficiency, providing strong support for complex post-translational modification research.", "normal"),
            ("", "normal"),
            ("### 4.5 Module Integration and Data Flow", "subheading"),
            ("", "normal"),
            ("Although the three analysis modules are deployed independently, they achieve organic integration at the data level. Through standardized data interfaces and unified data formats, the modules can achieve data sharing and result cross-referencing. For example, results from transcriptomics analysis can provide gene expression background for proteomics analysis, while proteomics results can provide protein expression baselines for post-translational modification analysis.", "normal"),
            ("", "normal"),
            ("The system architecture adopts microservice design concepts, with each module running on an independent port, achieving unified access through nginx reverse proxy. This architectural design not only improves system stability and scalability but also facilitates independent maintenance and upgrade of modules. Load balancing and fault tolerance mechanisms ensure system stability during high concurrent access, while modular deployment methods also facilitate resource allocation and performance optimization based on actual needs.", "normal"),
            ("", "normal"),
            ("Data security and user permission management are important considerations in system design. Each module implements user authentication and data access control, ensuring proper protection of sensitive biomedical data. Meanwhile, the system provides detailed operation logs and audit functions, facilitating tracking of data usage and ensuring compliance requirements.", "normal"),
            ("", "normal"),
            ("## Conclusion", "heading"),
            ("", "normal"),
            ("The ChatGIST project demonstrates how to organically combine modern web technologies with professional bioinformatics analysis tools, building a powerful, user-friendly intelligent analysis platform. The project's successful implementation proves the advantages of hybrid multi-stack architecture in complex application scenarios, maintaining both the professional features of each technology stack and achieving overall system coordination and optimization.", "normal"),
            ("", "normal"),
            ("The core value of the project lies in lowering the usage threshold of professional bioinformatics tools, allowing more researchers to conveniently conduct complex data analysis through modern user interfaces and intelligent data interpretation. The deep integration of AI functionality not only improves analysis efficiency but also provides professional scientific insights for researchers, promoting the development of biomedical research.", "normal"),
            ("", "normal"),
            ("From a technical implementation perspective, the modular design, containerized deployment, and microservice architecture adopted by the ChatGIST project provide valuable reference experiences for similar projects. This architectural design is not only applicable to the bioinformatics field but can also be extended to other complex application scenarios requiring integration of multiple professional tools and AI functions, providing new ideas and methods for modern web application development.", "normal"),
            ("", "normal")
        ]
        
        # 添加英文内容
        for text, style in english_content:
            # 添加新段落
            p = doc.add_paragraph(text)
            
            # 设置段落样式
            try:
                if style == "title":
                    # 标题样式
                    p.style = doc.styles['Title'] if 'Title' in [s.name for s in doc.styles] else doc.styles['Normal']
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
                    run.font.size = Pt(20)
                elif style == "heading":
                    # 一级标题样式
                    p.style = doc.styles['Heading 1'] if 'Heading 1' in [s.name for s in doc.styles] else doc.styles['Normal']
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
                    run.font.size = Pt(16)
                elif style == "subheading":
                    # 二级标题样式
                    p.style = doc.styles['Heading 2'] if 'Heading 2' in [s.name for s in doc.styles] else doc.styles['Normal']
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
                    run.font.size = Pt(14)
                else:
                    # 普通段落
                    p.style = doc.styles['Normal']
            except Exception as e:
                print(f"设置样式时出错: {e}")
                # 如果出错，使用默认样式
                p.style = doc.styles['Normal']
        
        # 保存英文版文档
        output_path = "ChatGIST_Project_Technical_Implementation_Details.docx"
        doc.save(output_path)
        print(f"英文版文档已成功创建并保存为: {output_path}")
        
        return True
        
    except Exception as e:
        print(f"创建英文版文档时发生错误: {str(e)}")
        return False

def main():
    """主函数"""
    print("ChatGIST项目技术文档英文版创建工具")
    print("=" * 50)
    
    # 创建英文版文档
    if create_english_version():
        print("英文版文档创建完成！")
        print("\n文档内容包括：")
        print("✅ 项目概述")
        print("✅ 系统架构")
        print("✅ 前端技术实现")
        print("✅ 后端API架构")
        print("✅ 数据库与存储设计")
        print("✅ AI功能实现")
        print("✅ 部署与运维")
        print("✅ 生物信息学分析模块")
        return 0
    else:
        print("英文版文档创建失败！")
        return 1

if __name__ == "__main__":
    sys.exit(main())
